#!/usr/bin/env python3
"""Convertit le cahier des charges Markdown en document Word."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color_hex: str):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_run(paragraph, text: str, bold=False, italic=False, code=False):
    if not text:
        return
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        if part.startswith('**') and part.endswith('**'):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run.text = part[1:-1]
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        else:
            run.text = part
            run.bold = bold
            run.italic = italic


def parse_table_row(line: str) -> list[str]:
    line = line.strip().strip('|')
    return [cell.strip() for cell in line.split('|')]


def is_separator_row(cells: list[str]) -> bool:
    return all(re.match(r'^:?-+:?$', c.replace(' ', '')) for c in cells if c)


def convert_md_to_docx(md_path: Path, docx_path: Path):
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style par défaut
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    lines = md_path.read_text(encoding='utf-8').splitlines()
    i = 0
    in_code_block = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []
    in_table = False
    list_buffer: list[tuple[str, int]] = []  # (text, level)

    def flush_list():
        nonlocal list_buffer
        for text, level in list_buffer:
            p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
            add_formatted_run(p, text)
        list_buffer = []

    def flush_table():
        nonlocal table_rows, in_table
        if not table_rows:
            in_table = False
            return
        cols = max(len(r) for r in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ri, row in enumerate(table_rows):
            for ci in range(cols):
                cell_text = row[ci] if ci < len(row) else ''
                cell = table.rows[ri].cells[ci]
                cell.text = ''
                p = cell.paragraphs[0]
                add_formatted_run(p, cell_text, bold=(ri == 0))
                p.paragraph_format.space_after = Pt(2)
                if ri == 0:
                    set_cell_shading(cell, '1F4E79')
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.bold = True
        doc.add_paragraph()
        table_rows = []
        in_table = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code block
        if stripped.startswith('```'):
            if in_code_block:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.3)
                code_lines = []
                in_code_block = False
            else:
                flush_list()
                flush_table()
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table
        if stripped.startswith('|') and stripped.endswith('|'):
            cells = parse_table_row(stripped)
            if is_separator_row(cells):
                i += 1
                continue
            flush_list()
            in_table = True
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            flush_table()

        # Empty line
        if not stripped:
            flush_list()
            i += 1
            continue

        # Horizontal rule
        if stripped == '---':
            flush_list()
            p = doc.add_paragraph()
            p.add_run('─' * 60)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if heading_match:
            flush_list()
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            if level == 1:
                h = doc.add_heading(text, level=0)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            elif level == 2:
                h = doc.add_heading(text, level=1)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            else:
                doc.add_heading(text, level=min(level, 3))
            i += 1
            continue

        # Blockquote
        if stripped.startswith('>'):
            flush_list()
            text = stripped.lstrip('> ').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Unordered list
        list_match = re.match(r'^(\s*)[-*]\s+(.+)$', line)
        if list_match:
            indent = len(list_match.group(1))
            level = 1 if indent >= 2 else 0
            list_buffer.append((list_match.group(2), level))
            i += 1
            continue

        # Ordered list
        ordered_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if ordered_match:
            flush_list()
            p = doc.add_paragraph(style='List Number')
            add_formatted_run(p, ordered_match.group(2))
            i += 1
            continue

        # Regular paragraph
        flush_list()
        p = doc.add_paragraph()
        add_formatted_run(p, stripped)
        i += 1

    flush_list()
    flush_table()

    # Pied de page
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('Document confidentiel — Prospera © 2026')
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(str(docx_path))
    print(f'Document Word créé : {docx_path}')


if __name__ == '__main__':
    base = Path(__file__).parent
    md_file = base / 'CAHIER_DES_CHARGES_PROSPERA_WEB.md'
    docx_file = base / 'CAHIER_DES_CHARGES_PROSPERA_WEB.docx'
    convert_md_to_docx(md_file, docx_file)
